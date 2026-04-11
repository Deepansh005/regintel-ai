import os
from pathlib import Path
from fastapi import APIRouter, Body, HTTPException
from fastapi.responses import FileResponse
from fpdf import FPDF
from docx import Document
from db.database import get_task
from app.services.pdf_service import extract_pdf_pages
from app.services.clause_extractor import extract_clauses_from_pages
from app.rag.vector_store import (
    list_collections,
    delete_collection,
    delete_all_collections,
    reset_chroma_db,
    get_chunk_details,
)

router = APIRouter()

EXPORTS_DIR = "exports"
os.makedirs(EXPORTS_DIR, exist_ok=True)

@router.get("/export/{task_id}/{format}")
def export_report(task_id: str, format: str):
    task = get_task(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Analysis report not found")

    result = task.get("result", {})
    filename = f"RegIntel_Report_{task_id[:8]}.{format}"
    filepath = os.path.join(EXPORTS_DIR, filename)

    try:
        if format == "pdf":
            generate_pdf(result, filepath)
        elif format == "docx":
            generate_docx(result, filepath)
        else:
            raise HTTPException(status_code=400, detail="Invalid format. Use 'pdf' or 'docx'")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate report: {str(e)}")

    return FileResponse(filepath, filename=filename, media_type='application/octet-stream')


@router.get("/chroma/collections")
def chroma_collections():
    collections = list_collections()
    return {
        "collections": [
            {"name": collection.name, "metadata": collection.metadata}
            for collection in collections
        ]
    }


@router.delete("/chroma/collections")
def chroma_delete_all_collections():
    delete_all_collections()
    return {"message": "All collections deleted successfully"}


@router.delete("/chroma/collections/{collection_name}")
def chroma_delete_collection(collection_name: str):
    try:
        delete_collection(collection_name)
        return {"message": f"Collection '{collection_name}' deleted successfully"}
    except Exception as exc:
        raise HTTPException(status_code=404, detail=f"Collection not found or could not be deleted: {str(exc)}")


@router.delete("/chroma/reset")
def chroma_reset():
    reset_chroma_db()
    return {"message": "ChromaDB folder deleted and fresh database initialized"}


@router.post("/chunks/details")
def chunk_details(payload: dict = Body(...)):
    chunk_ids = payload.get("chunk_ids") if isinstance(payload, dict) else None

    if not isinstance(chunk_ids, list) or not chunk_ids:
        raise HTTPException(status_code=400, detail="chunk_ids must be a non-empty list")

    return {
        "chunks": get_chunk_details(chunk_ids)
    }


@router.post("/debug/clause-extraction")
def debug_clause_extraction(payload: dict = Body(...)):
    file_path = payload.get("file_path") if isinstance(payload, dict) else None
    if not file_path:
        raise HTTPException(status_code=400, detail="file_path is required")

    normalized_path = str(Path(file_path))
    if not os.path.exists(normalized_path):
        raise HTTPException(status_code=404, detail=f"File not found: {normalized_path}")

    pages = extract_pdf_pages(normalized_path)
    clauses = extract_clauses_from_pages(pages)

    first_clause = clauses[0] if clauses else {}
    print(f"Clause debug endpoint: total_clauses={len(clauses)} first_clause={first_clause.get('content', '')[:500] if isinstance(first_clause, dict) else ''}")

    return {
        "total_clauses": len(clauses),
        "first_clause": first_clause,
        "clauses": clauses[:5],
    }

def generate_pdf(data, path):
    pdf = FPDF()
    pdf.add_page()
    
    # Title
    pdf.set_font("Helvetica", 'B', 24)
    pdf.set_text_color(124, 58, 237) # Violet color
    pdf.cell(0, 20, "RegIntel AI Compliance Report", ln=True, align='C')
    pdf.ln(10)

    # Helper to parse and format content
    def format_section(title, content):
        pdf.set_font("Helvetica", 'B', 16)
        pdf.set_text_color(31, 41, 55) # Slate 800
        pdf.cell(0, 10, title, ln=True)
        pdf.ln(2)
        
        pdf.set_font("Helvetica", '', 11)
        pdf.set_text_color(75, 85, 99) # Slate 600

        if isinstance(content, dict):
            # Special handling for different AI result keys
            if "changes" in content:
                for c in content["changes"]:
                    text = f"- [{c.get('type', 'Change').upper()}] {c.get('section', 'General')}: {c.get('summary', '')}"
                    pdf.multi_cell(0, 6, text)
                    pdf.ln(2)
            elif "gaps" in content:
                for g in content["gaps"]:
                    text = f"- ISSUE: {g.get('issue', '')}\n  RISK: {g.get('risk', 'Medium')}\n  REF: {g.get('policy_reference', 'N/A')}"
                    pdf.multi_cell(0, 6, text)
                    pdf.ln(2)
            elif "Impact" in content:
                imp = content["Impact"]
                pdf.multi_cell(0, 6, f"Departments: {', '.join(imp.get('departments', []))}")
                pdf.multi_cell(0, 6, f"Summary: {imp.get('summary', '')}")
            elif "actions" in content:
                for a in content["actions"]:
                    text = f"- ACTION: {a.get('step', '')}\n  OWNER: {a.get('owner', 'Compliance')}\n  TIMELINE: {a.get('timeline', 'Immediate')}"
                    pdf.multi_cell(0, 6, text)
                    pdf.ln(2)
            else:
                pdf.multi_cell(0, 6, str(content))
        else:
            pdf.multi_cell(0, 6, str(content))
        
        pdf.ln(8)

    format_section("Regulatory Changes", data.get("changes", {}))
    format_section("Compliance Gaps", data.get("compliance_gaps", {}))
    format_section("Strategic Impact", data.get("impact", {}))
    format_section("Recommended Actions", data.get("actions", {}))

    pdf.output(path)

def generate_docx(data, path):
    doc = Document()
    doc.add_heading('RegIntel AI Compliance Report', 0)

    def add_section(title, content):
        doc.add_heading(title, level=1)
        if isinstance(content, dict):
            if "changes" in content:
                for c in content["changes"]:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f"[{c.get('type', 'Change').upper()}] {c.get('section', 'General')}: ").bold = True
                    p.add_run(c.get('summary', ''))
            elif "gaps" in content:
                for g in content["gaps"]:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f"ISSUE: {g.get('issue', '')}").bold = True
                    p.add_run(f"\nRISK: {g.get('risk', 'Medium')}\nREF: {g.get('policy_reference', 'N/A')}")
            elif "Impact" in content:
                imp = content["Impact"]
                doc.add_paragraph(f"Departments Involved: {', '.join(imp.get('departments', []))}")
                doc.add_paragraph(imp.get('summary', ''))
            elif "actions" in content:
                for a in content["actions"]:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f"ACTION: {a.get('step', '')}").bold = True
                    p.add_run(f"\nOWNER: {a.get('owner', 'N/A')} | TIMELINE: {a.get('timeline', 'Immediate')}")
            else:
                doc.add_paragraph(str(content))
        else:
            doc.add_paragraph(str(content))

    add_section("Regulatory Changes", data.get("changes", {}))
    add_section("Compliance Gaps", data.get("compliance_gaps", {}))
    add_section("Strategic Impact", data.get("impact", {}))
    add_section("Recommended Actions", data.get("actions", {}))

    doc.save(path)
